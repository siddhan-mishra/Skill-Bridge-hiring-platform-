import re
import os
import io
import time
import spacy
import json
import numpy as np
from spacy.matcher import PhraseMatcher
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from skillNer.skill_extractor_class import SkillExtractor
import pdfplumber
from google import genai as genai_new
from dotenv import load_dotenv

load_dotenv()

try:
    from skillNer.general_params import SKILL_DB
except ImportError:
    from skillNer.utils import load_skill_db
    SKILL_DB = load_skill_db()

_sbert_model = None

def get_sbert():
    global _sbert_model
    if _sbert_model is None:
        from sentence_transformers import SentenceTransformer
        print("[sbert] loading all-MiniLM-L6-v2 ...")
        _sbert_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("[sbert] model loaded")
    return _sbert_model

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

nlp = spacy.load("en_core_web_sm")
skill_extractor = SkillExtractor(nlp, SKILL_DB, PhraseMatcher)

# ── Gemini client ────────────────────────────────────────────────────
_gemini_key = os.environ.get("GEMINI_API_KEY")
if not _gemini_key:
    raise RuntimeError("GEMINI_API_KEY not set in nlp-service/.env")
gemini = genai_new.Client(api_key=_gemini_key)

GEMINI_MODELS = [
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
]

# ── Groq client (free fallback) ─────────────────────────────────────────
# Get a FREE key at: https://console.groq.com  (14,400 req/day free)
# Add GROQ_API_KEY=your_key to nlp-service/.env
_groq_key = os.environ.get("GROQ_API_KEY", "")
_groq_client = None
if _groq_key:
    try:
        from groq import Groq
        _groq_client = Groq(api_key=_groq_key)
        print("[groq] client initialised ✔")
    except ImportError:
        print("[groq] groq package not installed — run: pip install groq")
else:
    print("[groq] GROQ_API_KEY not set — Groq fallback disabled")

GROQ_MODELS = [
    "llama-3.3-70b-versatile",   # best quality, free tier
    "llama3-70b-8192",           # fast fallback
    "llama3-8b-8192",            # smallest, always available
]

BONUS_EXPANSIONS = {
    'mern':       ['mongodb', 'express', 'react', 'node'],
    'mean':       ['mongodb', 'express', 'angular', 'node'],
    'mevn':       ['mongodb', 'express', 'vue', 'node'],
    'lamp':       ['linux', 'apache', 'mysql', 'php'],
    'fullstack':  ['html', 'css', 'javascript', 'react', 'node', 'mongodb'],
    'full stack': ['html', 'css', 'javascript', 'react', 'node', 'mongodb'],
    'devops':     ['docker', 'kubernetes', 'ci/cd', 'linux', 'git'],
    'data science': ['python', 'pandas', 'numpy', 'machine learning', 'sql'],
    'ml':         ['python', 'machine learning', 'scikit-learn', 'pandas', 'numpy'],
    'ai/ml':      ['python', 'machine learning', 'deep learning', 'tensorflow', 'pytorch'],
}


class TextIn(BaseModel):
    text: str

class SkillsOut(BaseModel):
    skills: List[str]

class SkillSuggestIn(BaseModel):
    title: str = ""
    description: str = ""

class MatchRequest(BaseModel):
    profileText:   str
    jobText:       str
    profileSkills: List[str] = []
    jobSkills:     List[str] = []
    profileYears:  Optional[int] = 0
    requiredYears: Optional[int] = 0
    educationMatch: Optional[float] = 0.5

class ResumeRequest(BaseModel):
    name:          str
    email:         Optional[str] = ""
    phone:         Optional[str] = ""
    location:      Optional[str] = ""
    linkedinUrl:   Optional[str] = ""
    githubUrl:     Optional[str] = ""
    portfolioUrl:  Optional[str] = ""
    headline:      Optional[str] = ""
    summary:       Optional[str] = ""
    skills:        Optional[List[str]] = []
    techStack:     Optional[List[str]] = []
    softSkills:    Optional[List[str]] = []
    workHistory:   Optional[List[Dict[str, Any]]] = []
    education:     Optional[List[Dict[str, Any]]] = []
    projects:      Optional[List[Dict[str, Any]]] = []
    certifications: Optional[List[Dict[str, Any]]] = []
    languages:     Optional[List[str]] = []
    yearsOfExp:    Optional[int] = 0


@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "gemini_models": GEMINI_MODELS,
        "groq_enabled": _groq_client is not None,
        "groq_models": GROQ_MODELS if _groq_client else [],
    }


# ────────────────────────────────────────────────────────────────────────────────
def call_groq(prompt: str) -> str:
    """Call Groq LLaMA — free tier, 14,400 req/day."""
    if not _groq_client:
        raise RuntimeError("Groq not configured (GROQ_API_KEY missing or groq package not installed)")
    last_err = None
    for model in GROQ_MODELS:
        try:
            print(f"[groq] trying model={model}")
            resp = _groq_client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=4096,
            )
            text = resp.choices[0].message.content.strip()
            print(f"[groq] success model={model}")
            return text
        except Exception as ex:
            msg = str(ex)
            last_err = ex
            if "429" in msg or "rate_limit" in msg.lower():
                print(f"[groq] 429 on {model}, waiting 10s")
                time.sleep(10)
                continue
            else:
                print(f"[groq] error on {model}: {ex}")
                continue
    raise last_err or RuntimeError("All Groq models exhausted")


def call_ai(prompt: str) -> str:
    """
    Primary: Gemini (fast, high quality).
    Fallback: Groq/LLaMA (free, 14.4k req/day) when Gemini quota exhausted.
    """
    gemini_exhausted = False
    last_gemini_err = None

    for attempt, model in enumerate(GEMINI_MODELS):
        try:
            print(f"[gemini] trying model={model}")
            resp = gemini.models.generate_content(model=model, contents=prompt)
            print(f"[gemini] success model={model}")
            return resp.text.strip()
        except Exception as ex:
            msg = str(ex)
            last_gemini_err = ex
            if "429" in msg or "RESOURCE_EXHAUSTED" in msg:
                # Check if daily quota is gone (limit: 0) — no point waiting
                if "limit: 0" in msg or "GenerateRequestsPerDay" in msg:
                    print(f"[gemini] daily quota EXHAUSTED on {model} — switching to Groq")
                    gemini_exhausted = True
                    break
                delay = 10 + attempt * 5
                m = re.search(r'retryDelay.*?(\d+)s', msg)
                if m:
                    delay = min(int(m.group(1)) + attempt * 3, 30)
                print(f"[gemini] 429 on {model}, waiting {delay}s then trying next model")
                time.sleep(delay)
                continue
            elif "404" in msg or "NOT_FOUND" in msg:
                print(f"[gemini] 404 on {model} — skipping")
                continue
            else:
                print(f"[gemini] error on {model}: {ex}")
                continue

    # ── Gemini failed — try Groq ──
    if _groq_client:
        print("[ai] Gemini unavailable — falling back to Groq/LLaMA")
        return call_groq(prompt)

    raise last_gemini_err or RuntimeError("All AI models exhausted — add GROQ_API_KEY for a free fallback")


# Keep old name for internal use
def call_gemini(prompt: str) -> str:
    return call_ai(prompt)


def strip_fences(raw: str) -> str:
    if "```" not in raw:
        return raw
    parts = raw.split("```")
    for part in parts:
        p = part.strip()
        if p.startswith("json"):
            p = p[4:].strip()
        if p.startswith("{") or p.startswith("["):
            return p
    return raw


# ────────────────────────────────────────────────────────────────────────────────
@app.post("/extract-skills", response_model=SkillsOut)
def extract_skills(payload: TextIn):
    doc = skill_extractor.annotate(payload.text)
    results = doc.get("results", {})
    skills = set()
    for key in ["full_matches", "ngram_scored_filtered", "ngram_scored", "keyword_scored"]:
        for item in (results.get(key) or []):
            if isinstance(item, dict):
                name = item.get("skill_name") or item.get("doc_node_value")
                if name:
                    skills.add(name.strip())
    return {"skills": sorted(skills)}


@app.post("/suggest-skills")
async def suggest_skills(payload: SkillSuggestIn):
    if not payload.title and not payload.description:
        return {"skills": []}

    PROMPT = f"""You are an expert technical recruiter with deep knowledge of the IT and software industry.

Given the job title and description below, return a JSON array of the most relevant technical and soft skills
required for this role. Follow these rules strictly:

1. Return ONLY a raw JSON array of strings. No markdown, no explanation.
2. Expand shorthand: "MERN" becomes ["MongoDB", "Express.js", "React", "Node.js"] as separate items.
3. Include both hard skills (technologies, frameworks, languages) and soft skills.
4. Normalize capitalization: "JavaScript" not "javascript", "Node.js" not "nodejs".
5. Return 15-25 skills maximum, ordered by importance.
6. Do not include generic terms like "Computer Science" or "Software Development".

Job Title: {payload.title}
Job Description: {payload.description[:3000]}

Return ONLY the JSON array:"""

    try:
        raw = call_ai(PROMPT)
        raw = strip_fences(raw)
        parsed = json.loads(raw)
        skills = parsed if isinstance(parsed, list) else parsed.get("skills", [])
        skills = [str(s).strip()[:50] for s in skills if s]
        return {"skills": skills}
    except Exception as ex:
        print(f"[suggest-skills] error: {ex}")
        return {"skills": [], "error": str(ex)}


@app.post("/compute-match")
async def compute_match(req: MatchRequest):
    try:
        sbert = get_sbert()
        from sklearn.metrics.pairwise import cosine_similarity

        profile_emb = sbert.encode([req.profileText], normalize_embeddings=True)
        job_emb     = sbert.encode([req.jobText],     normalize_embeddings=True)
        sbert_score = float(cosine_similarity(profile_emb, job_emb)[0][0])
        sbert_score = max(0.0, sbert_score)

        def expand_skills(skill_list):
            expanded = set()
            for s in skill_list:
                sl = s.strip().lower()
                expanded.add(sl)
                if sl in BONUS_EXPANSIONS:
                    expanded.update(BONUS_EXPANSIONS[sl])
            return expanded

        profile_set = expand_skills(req.profileSkills)
        job_set     = set(s.strip().lower() for s in req.jobSkills)

        if not job_set:
            skill_score = 1.0
        else:
            matched = sum(1 for js in job_set if js in profile_set or
                          any(_fuzzy(js, ps) for ps in profile_set))
            skill_score = matched / len(job_set)

        matched_skills = [js for js in job_set if js in profile_set or
                           any(_fuzzy(js, ps) for ps in profile_set)]
        missing_skills = [js for js in job_set if js not in matched_skills]
        extra_skills   = [ps for ps in profile_set
                          if ps not in job_set and
                          not any(_fuzzy(ps, js) for js in job_set)]

        exp_years = req.profileYears or 0
        req_years = req.requiredYears or 0
        if req_years == 0:
            exp_score = 1.0
        elif exp_years >= req_years:
            exp_score = 1.0
        elif exp_years >= req_years * 0.7:
            exp_score = 0.7
        else:
            exp_score = max(0.2, exp_years / req_years)

        edu_score = req.educationMatch
        structured_score = (exp_score * 0.6) + (edu_score * 0.4)

        final_raw = (0.40 * sbert_score + 0.40 * skill_score + 0.20 * structured_score)
        final_pct = round(final_raw * 100)

        print(f"[compute-match] sbert={sbert_score:.3f} skill={skill_score:.3f} struct={structured_score:.3f} → {final_pct}%")

        return {
            "score": final_pct,
            "breakdown": {
                "semanticScore": round(sbert_score * 100),
                "skillScore":    round(skill_score * 100),
                "structScore":   round(structured_score * 100),
            },
            "matchedSkills": matched_skills,
            "missingSkills": missing_skills,
            "extraSkills":   list(extra_skills)[:10],
        }
    except Exception as ex:
        print(f"[compute-match] error: {ex}")
        return {"score": -1, "error": str(ex)}


@app.post("/generate-resume")
async def generate_resume(req: ResumeRequest):
    profile_json = req.model_dump()
    PROMPT = f"""You are a Harvard Career Services resume writer. Given the raw profile data below,
write polished resume content following Harvard format rules:

RULES:
1. Return ONLY a raw JSON object — no markdown, no backticks.
2. summary: 2-3 concise sentences, 3rd person, professional tone.
3. workHistory: for each job, rewrite achievements as 2-4 action-verb led bullet strings.
   Use strong verbs: Built, Engineered, Reduced, Increased, Led, Designed, Deployed.
   Add measurable impact where possible: "Reduced API latency by 40%".
4. skills_line: single comma-separated string of top 12 technical skills.
5. projects: for each project, write 1-2 bullet impact strings.
6. If a field is empty/missing, return empty string or empty array — never null.
7. Do NOT invent companies, dates, or facts not in the profile.

EXACT OUTPUT FORMAT:
{{
  "summary": "Experienced full-stack developer...",
  "skills_line": "JavaScript, React, Node.js, Python, MongoDB, Docker, AWS, Git",
  "workHistory": [
    {{
      "role": "Software Engineer",
      "company": "Google",
      "startDate": "Jan 2021",
      "endDate": "Present",
      "bullets": ["Built REST APIs handling 10M+ daily requests", "Reduced deployment time by 60%"]
    }}
  ],
  "projects": [
    {{
      "title": "SkillBridge",
      "technologies": "React, Node.js, MongoDB",
      "bullets": ["Engineered AI matching engine with 85% accuracy"]
    }}
  ]
}}

RAW PROFILE:
{json.dumps(profile_json, indent=2)[:6000]}

Return ONLY the JSON:"""

    ai_content = {}
    try:
        raw = call_ai(PROMPT)
        raw = strip_fences(raw)
        ai_content = json.loads(raw)
        print(f"[generate-resume] AI OK for {req.name}")
    except Exception as ex:
        print(f"[generate-resume] AI error, using raw profile data: {ex}")
        ai_content = {
            "summary": req.summary or "",
            "skills_line": ", ".join((req.skills or [])[:12]),
            "workHistory": [
                {"role": w.get("role",""), "company": w.get("company",""),
                 "startDate": w.get("startDate",""), "endDate": w.get("endDate",""),
                 "bullets": w.get("achievements", [])}
                for w in (req.workHistory or [])
            ],
            "projects": [
                {"title": p.get("title",""), "technologies": ", ".join(p.get("technologies",[]) if isinstance(p.get("technologies"), list) else [p.get("technologies","")]),
                 "bullets": [p.get("description","")]}
                for p in (req.projects or [])
            ],
        }

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    def set_font(run, size=11, bold=False, color=None):
        run.font.name  = "Garamond"
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_hrule(doc):
        p  = doc.add_paragraph()
        pr = p._p.get_or_add_pPr()
        pb = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'),  '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pb.append(bottom)
        pr.append(pb)
        p.paragraph_format.space_after = Pt(4)

    def add_section_heading(doc, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title.upper())
        set_font(run, size=11, bold=True)
        add_hrule(doc)

    def add_bullet(doc, text, indent=0.25):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Inches(indent)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        set_font(run, size=10.5)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(req.name or "")
    set_font(name_run, size=16, bold=True)

    contact_parts = [p for p in [
        req.phone, req.email, req.location,
        req.linkedinUrl, req.githubUrl, req.portfolioUrl
    ] if p]
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    contact_run = contact_p.add_run("  |  ".join(contact_parts))
    set_font(contact_run, size=10)

    summary_text = ai_content.get("summary") or req.summary or ""
    if summary_text:
        add_section_heading(doc, "Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(summary_text)
        set_font(run, size=10.5)

    skills_line = ai_content.get("skills_line") or ", ".join((req.skills or [])[:12])
    if skills_line:
        add_section_heading(doc, "Technical Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(skills_line)
        set_font(run, size=10.5)

    ai_work = ai_content.get("workHistory") or []
    if ai_work:
        add_section_heading(doc, "Experience")
        for job in ai_work:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            role_run = p.add_run(f"{job.get('role','')}, ")
            set_font(role_run, size=11, bold=True)
            company_run = p.add_run(job.get('company', ''))
            set_font(company_run, size=11)
            dates = f"{job.get('startDate','')} – {job.get('endDate','Present')}"
            dates_run = p.add_run(f"  |  {dates}")
            set_font(dates_run, size=10.5, color=(100, 100, 100))
            for bullet in (job.get("bullets") or []):
                if bullet:
                    add_bullet(doc, bullet)

    if req.education:
        add_section_heading(doc, "Education")
        for edu in req.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            deg_run = p.add_run(f"{edu.get('degree','')}, ")
            set_font(deg_run, size=11, bold=True)
            inst_run = p.add_run(edu.get('institute', ''))
            set_font(inst_run, size=11)
            yr = edu.get('year', '')
            gpa = edu.get('gpa', '')
            meta = " | ".join(filter(None, [yr, f"GPA: {gpa}" if gpa else ""]))
            if meta:
                meta_run = p.add_run(f"  |  {meta}")
                set_font(meta_run, size=10.5, color=(100, 100, 100))

    ai_projects = ai_content.get("projects") or []
    if ai_projects:
        add_section_heading(doc, "Projects")
        for proj in ai_projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            title_run = p.add_run(proj.get('title', ''))
            set_font(title_run, size=11, bold=True)
            tech = proj.get('technologies', '')
            if tech:
                tech_run = p.add_run(f"  —  {tech}")
                set_font(tech_run, size=10.5, color=(100, 100, 100))
            for bullet in (proj.get("bullets") or []):
                if bullet:
                    add_bullet(doc, bullet)

    if req.certifications:
        add_section_heading(doc, "Certifications")
        for cert in req.certifications:
            name = cert.get('name', '')
            issuer = cert.get('issuer', '')
            year   = cert.get('year', '')
            line   = " | ".join(filter(None, [name, issuer, year]))
            if line:
                add_bullet(doc, line)

    if req.languages:
        add_section_heading(doc, "Languages")
        p = doc.add_paragraph()
        run = p.add_run(", ".join(req.languages))
        set_font(run, size=10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r'[^\w\s-]', '', req.name or 'Resume').replace(' ', '_')
    filename  = f"{safe_name}_Resume.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _fuzzy(a: str, b: str) -> bool:
    if a == b: return True
    if len(a) <= 3 or len(b) <= 3: return False
    if abs(len(a) - len(b)) > 2: return False
    m, n = len(a), len(b)
    dp = list(range(n + 1))
    for i in range(1, m + 1):
        prev = dp[:]
        dp[0] = i
        for j in range(1, n + 1):
            dp[j] = prev[j - 1] if a[i-1] == b[j-1] else 1 + min(prev[j], dp[j-1], prev[j-1])
    return dp[n] <= 1


def pdf_to_text(data: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()


def rx_email(t):     m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', t);                 return m.group(0) if m else None
def rx_phone(t):     m = re.search(r'(\+?\d[\d\s\-().]{7,}\d)', t);              return m.group(0).strip() if m else None
def rx_linkedin(t):  m = re.search(r'linkedin\.com/in/[\w\-]+', t, re.I);         return "https://" + m.group(0) if m else None
def rx_github(t):    m = re.search(r'github\.com/[\w\-]+', t, re.I);              return "https://" + m.group(0) if m else None
def rx_portfolio(t):
    hits = re.findall(r'https?://(?!.*linkedin)(?!.*github)[\w\-\.]+\.[a-z]{2,}[\S]*', t, re.I)
    return hits[0] if hits else None


@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    content = await file.read()
    text = pdf_to_text(content)

    if not text:
        return {"error": "No text extracted — PDF may be image-based (scanned)."}

    email     = rx_email(text)
    phone     = rx_phone(text)
    linkedin  = rx_linkedin(text)
    github    = rx_github(text)
    portfolio = rx_portfolio(text)

    PROMPT = f"""You are an expert resume parser. Extract structured data from the resume below.

RULES (follow strictly):
1. Return ONLY a raw JSON object. No markdown, no triple backticks, no explanation.
2. Use null (JSON null) for truly missing fields.
3. Arrays must be [] if nothing found — never null.
4. yearsOfExp = integer (e.g. 3), not a string.
5. skills = programming languages, frameworks, libraries, technical concepts.
6. tools = software tools, platforms, IDEs, cloud services (Git, Docker, AWS, VS Code).
7. achievements = list of bullet strings from job description.
8. education year = 4-digit graduation year string e.g. "2024".
9. workHistory endDate = "Present" if currently working there.

EXACT OUTPUT FORMAT:
{{
  "fullName": "Jane Smith",
  "headline": "Full Stack Developer with 3 years experience",
  "summary": "Experienced developer specializing in MERN stack...",
  "currentTitle": "Software Engineer",
  "currentCompany": "Google",
  "yearsOfExp": 3,
  "skills": ["JavaScript", "React", "Node.js", "Python", "MongoDB"],
  "tools": ["Git", "Docker", "VS Code", "Postman", "AWS"],
  "workHistory": [
    {{
      "role": "Software Engineer",
      "company": "Google",
      "startDate": "Jan 2021",
      "endDate": "Present",
      "achievements": [
        "Built REST APIs serving 10M+ requests/day",
        "Reduced page load time by 40%"
      ]
    }}
  ],
  "education": [
    {{
      "degree": "B.Tech Computer Science",
      "institute": "KIIT University",
      "year": "2021",
      "gpa": "8.5"
    }}
  ],
  "certifications": [
    {{
      "name": "AWS Solutions Architect",
      "issuer": "Amazon",
      "year": "2022",
      "link": ""
    }}
  ],
  "languages": ["English", "Hindi"],
  "projects": [
    {{
      "title": "SkillBridge",
      "description": "AI-powered job matching platform",
      "technologies": ["React", "Node.js", "MongoDB"],
      "link": "https://github.com/user/skillbridge"
    }}
  ]
}}

RESUME TEXT:
---
{text[:9000]}
---"""

    parsed = {}
    try:
        raw = call_ai(PROMPT)
        raw = strip_fences(raw)
        parsed = json.loads(raw)
    except Exception as ex:
        print(f"[parse-resume] error: {ex}")

    return {
        "fullName":       parsed.get("fullName"),
        "email":          email,
        "phone":          phone,
        "linkedinUrl":    linkedin,
        "githubUrl":      github,
        "portfolioUrl":   portfolio,
        "headline":       parsed.get("headline"),
        "summary":        parsed.get("summary"),
        "currentTitle":   parsed.get("currentTitle"),
        "currentCompany": parsed.get("currentCompany"),
        "yearsOfExp":     parsed.get("yearsOfExp"),
        "skills":         parsed.get("skills")         or [],
        "tools":          parsed.get("tools")          or [],
        "workHistory":    parsed.get("workHistory")    or [],
        "education":      parsed.get("education")      or [],
        "certifications": parsed.get("certifications") or [],
        "languages":      parsed.get("languages")      or [],
        "projects":       parsed.get("projects")       or [],
    }
